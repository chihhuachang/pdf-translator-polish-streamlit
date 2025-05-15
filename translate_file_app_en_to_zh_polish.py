import streamlit as st
import google.generativeai as genai
import os
import sys
import io
import time

# Import libraries for file reading AND writing docx
import docx
from docx import Document
import PyPDF2
from PyPDF2.errors import PdfReadError

# --- 常數定義 (保持不變) ---
MAX_CHARS_PER_CHUNK = 2500
API_CALL_DELAY = 2 # 秒

# --- 說明文件 ---
# 功能：上傳英文文件，分塊翻譯成繁體中文，合併後，
#       再對合併結果進行 **潤飾與校對**，並分別提供初步翻譯和潤飾後翻譯的 Docx 下載。
# (其餘說明與之前版本相同)
#
# 如何執行：
# 1. 安裝函式庫: pip install streamlit google-generativeai python-docx PyPDF2
# 2. 將此程式碼儲存為 `translate_file_app_en_to_zh_polish.py`。
# 3. 在終端機中執行： streamlit run translate_file_app_en_to_zh_polish.py
# ---

# --- API 金鑰設定 (保持不變) ---
api_key = os.getenv("GOOGLE_API_KEY")
# ...(省略 API Key 檢查與設定程式碼)...
if not api_key: st.error("..."); st.stop()
try: genai.configure(api_key=api_key)
except Exception as e: st.error(f"...: {e}"); st.stop()

# --- 從檔案提取文字的函式 (保持不變) ---
def extract_text_from_file(uploaded_file):
    """ (此函式邏輯與上一版本完全相同) """
    extracted_text = ""
    # ...(省略檔案讀取和錯誤處理邏輯)...
    try:
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        if file_extension == ".txt":
            try: extracted_text = uploaded_file.getvalue().decode("utf-8")
            except UnicodeDecodeError:
                st.warning("嘗試 UTF-8 解碼失敗，嘗試 Big5...")
                try: extracted_text = uploaded_file.getvalue().decode("big5", errors='ignore')
                except Exception as e_enc:
                     st.error(f"嘗試 Big5 解碼也失敗: {e_enc}。使用忽略錯誤的 UTF-8。")
                     extracted_text = uploaded_file.getvalue().decode("utf-8", errors='ignore')
            st.info(f"成功讀取 .txt 檔案: {uploaded_file.name}")
        elif file_extension == ".docx":
            document = docx.Document(uploaded_file)
            extracted_text = '\n'.join([para.text for para in document.paragraphs])
            st.info(f"成功讀取 .docx 檔案: {uploaded_file.name}")
        elif file_extension == ".pdf":
            try:
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                if pdf_reader.is_encrypted: st.error("錯誤：PDF 文件已加密。"); return None
                full_text = [page.extract_text() for i, page in enumerate(pdf_reader.pages) if page.extract_text() or st.warning(f"讀取 PDF 第 {i+1} 頁時未提取到文字或發生錯誤。", icon="⚠️")]
                extracted_text = '\n'.join(filter(None, full_text))
                if not extracted_text.strip(): st.warning(f"無法從 PDF '{uploaded_file.name}' 提取任何文字。")
                else: st.info(f"成功讀取 .pdf 檔案: {uploaded_file.name}")
            except PdfReadError as pdf_err: st.error(f"PyPDF2 讀取錯誤: {pdf_err}"); return None
        else: st.error(f"錯誤：不支援的檔案類型 '{file_extension}'。"); return None
        return extracted_text.strip()
    except Exception as e: st.error(f"讀取或解析檔案 '{uploaded_file.name}' 時發生錯誤: {e}"); return None


# --- 文本分塊函式 (保持不變) ---
def split_text_into_chunks(text, max_chars=MAX_CHARS_PER_CHUNK):
    """ (此函式邏輯與上一版本完全相同) """
    chunks = []
    # ...(省略分塊邏輯)...
    paragraphs = text.split('\n\n')
    current_chunk = ""
    for paragraph in paragraphs:
        if len(paragraph) > max_chars:
            start = 0
            while start < len(paragraph):
                end = min(start + max_chars, len(paragraph))
                long_para_chunk = paragraph[start:end]
                if len(current_chunk) + len(long_para_chunk) + 2 <= max_chars:
                     if current_chunk: current_chunk += "\n\n" + long_para_chunk
                     else: current_chunk = long_para_chunk
                else:
                    if current_chunk: chunks.append(current_chunk)
                    current_chunk = long_para_chunk
                    if len(current_chunk) >= max_chars: chunks.append(current_chunk); current_chunk = ""
                start = end
        else:
            if len(current_chunk) + len(paragraph) + 2 <= max_chars:
                if current_chunk: current_chunk += "\n\n" + paragraph
                else: current_chunk = paragraph
            else:
                chunks.append(current_chunk)
                current_chunk = paragraph
    if current_chunk: chunks.append(current_chunk)
    return chunks

# --- 翻譯函式 (保持使用固定的詳細提示詞和流式傳輸) ---
def translate_text(text_to_translate, target_language="繁體中文"):
    """ (此函式邏輯與上一版本完全相同，使用你提供的固定提示詞) """
    if not text_to_translate: return None
    fixed_instruction_prompt = """
Please act as a professional Chinese translator and you will actually follow the steps below to produce a natural and professional Chinese translation.
1. Carefully read and fully understand the original text, ensuring thorough comprehension without haste.
2. Carefully think and consider how you would share the content you just read with your imagined audience in Chinese.
3. Start to translate the text by writting down the proposed sharing content you just had with your imagined audience using traditional Chinese characters. Avoid translating word-for-word; aim for a comfortable, natural, and smooth manner of expression.
""".strip()
    full_prompt = f"{fixed_instruction_prompt}\n\n{text_to_translate}"
    model = genai.GenerativeModel('gemini-1.5-flash-latest')
    try:
        response = model.generate_content(full_prompt, stream=True)
        full_translated_text = "".join(chunk.text for chunk in response if hasattr(chunk, 'text') and chunk.text)
        return full_translated_text.strip()
    except Exception as e:
        st.error(f"翻譯塊 '{text_to_translate[:30]}...' 時 API 呼叫或流式處理過程中發生錯誤: {e}")
        return f"[[翻譯錯誤於塊: {text_to_translate[:30]}... - {e}]]"


# --- 新增：潤飾翻譯文本函式 ---
def polish_translation(raw_translated_text, target_language="繁體中文"):
    """
    使用 Gemini API 對初步翻譯的文本進行分析、校對、去除贅字與語法潤飾。
    """
    if not raw_translated_text or not raw_translated_text.strip():
        st.warning("沒有可潤飾的初步翻譯內容。")
        return None

    # 專為潤飾設計的提示詞
    polishing_prompt_instruction = f"""
以下是一段由英文逐塊翻譯再合併而成的{target_language}文本。
由於是分塊翻譯後合併，可能存在以下問題：
1.  語句之間銜接不夠自然流暢。
2.  可能出現因分塊導致的語意中斷或重複。
3.  可能包含模型因分塊不完整而產生的提示性語句（例如："The text cuts off here", "以下篇幅過長" 等類似訊息）。
4.  整體風格可能不夠統一。

請你扮演一位資深的中文編輯，執行以下任務：
1.  **分析並理解**提供的文本內容。
2.  **校對語法錯誤**，修正任何不正確的表達。
3.  **去除贅字和重複**，使語言更精煉。
4.  **潤飾語句**，確保整篇文本語氣連貫、表達自然、流暢易讀，符合專業的{target_language}書寫風格。
5.  **移除或修正**任何由分塊翻譯產生的不必要提示性語句或中斷標記。
6.  確保最終輸出的文本意思忠於原文（雖然你看不到原始英文，但要基於提供的中文譯文使其更完美）。
7.  請直接輸出潤飾後的完整{target_language}文本，不要包含任何額外的解釋或開頭語。

待潤飾的{target_language}文本如下：
""".strip()

    full_polishing_prompt = f"{polishing_prompt_instruction}\n\n{raw_translated_text}"

    # 為潤飾任務選擇模型，Pro 系列可能更佳，但 Flash 也能嘗試
    # model = genai.GenerativeModel('gemini-1.5-pro-latest')
    model = genai.GenerativeModel('gemini-1.5-flash-latest') # 保持與翻譯一致，或可升級

    st.info("正在將初步翻譯結果傳送給模型進行潤飾與校對...")

    try:
        response = model.generate_content(full_polishing_prompt, stream=True) # 同樣使用流式
        polished_text = "".join(chunk.text for chunk in response if hasattr(chunk, 'text') and chunk.text)
        return polished_text.strip()
    except Exception as e:
        st.error(f"潤飾翻譯時 API 呼叫或流式處理過程中發生錯誤: {e}")
        return f"[[潤飾錯誤: {e}]]"


# --- 從文字建立 Docx 檔案函式 (保持不變) ---
def create_docx_from_text(text_content, base_filename, suffix=""):
    """ (此函式微調，加入檔名後綴) """
    try:
        document = Document()
        for paragraph in text_content.split('\n'): document.add_paragraph(paragraph)
        docx_buffer = io.BytesIO()
        document.save(docx_buffer)
        docx_buffer.seek(0)
        docx_filename = f"{base_filename}{suffix}.docx" # 加入後綴
        return {'name': docx_filename, 'data': docx_buffer}, None
    except Exception as e:
        error_message = f"建立 Docx 檔案 '{base_filename}{suffix}.docx' 時發生錯誤: {e}"
        st.error(error_message)
        return {'name': None, 'data': None}, error_message


# --- Streamlit 應用程式介面 (調整輸出區) ---
st.set_page_config(page_title="文件翻譯+潤飾 (英->繁中)", layout="wide")
st.title("📝 Gemini 文件翻譯與潤飾 (英文 ➔ 繁體中文)")
st.caption(f"上傳文件，分塊翻譯，合併後再進行AI潤飾，並提供 Docx 下載")

# ...(省略 col1, uploaded_file, base_filename, translate_button 定義)...
col1, col2 = st.columns([2,3]) # 調整欄位寬度比例
uploaded_file = None
base_filename = None
with col1:
    st.subheader("步驟 1: 上傳英文文件")
    uploaded_file = st.file_uploader(
        "選擇要翻譯的英文文件 (.txt, .docx, .pdf)", type=['txt', 'docx', 'pdf'], key="file_uploader"
    )
    if uploaded_file is not None:
        st.markdown(f"**已上傳檔案:** `{uploaded_file.name}` (`{uploaded_file.type}`)")
        base_filename = os.path.splitext(uploaded_file.name)[0]
    translate_button = st.button("開始翻譯與潤飾", key="translate_btn", disabled=uploaded_file is None)

# --- 翻譯結果與潤飾結果輸出區 ---
with col2:
    st.subheader("初步翻譯結果 (繁體中文)")
    raw_result_placeholder = st.empty()
    raw_result_placeholder.text_area(
        label="raw_translation_label", value="初步翻譯結果將顯示於此...", height=250,
        key="raw_result_text_area", disabled=True, label_visibility="collapsed"
    )
    raw_download_placeholder = st.empty()

    st.markdown("---") # 分隔線

    st.subheader("AI 潤飾後結果 (繁體中文)")
    polished_result_placeholder = st.empty()
    polished_result_placeholder.text_area(
        label="polished_translation_label", value="AI 潤飾後的翻譯結果將顯示於此...", height=250,
        key="polished_result_text_area", disabled=True, label_visibility="collapsed"
    )
    polished_download_placeholder = st.empty()

progress_placeholder = st.empty() # 將進度條移到按鈕下方全域顯示

# --- 主執行邏輯 (加入潤飾步驟) ---
if translate_button:
    if uploaded_file is None:
        st.warning("請先上傳一個文件。")
    else:
        # 清空先前結果
        raw_download_placeholder.empty()
        polished_download_placeholder.empty()
        progress_placeholder.empty()
        raw_result_placeholder.text_area(label="raw_translation_label", value="處理中...", height=250, key="raw_processing", disabled=True, label_visibility="collapsed")
        polished_result_placeholder.text_area(label="polished_translation_label", value="等待初步翻譯完成...", height=250, key="polished_waiting", disabled=True, label_visibility="collapsed")

        with st.spinner(f"正在讀取檔案 '{uploaded_file.name}'..."):
            extracted_text = extract_text_from_file(uploaded_file)

        if extracted_text is not None and extracted_text.strip():
            st.success("成功從文件中提取英文文字！")

            # 步驟 2: 分塊
            with st.spinner("正在將文本分割成處理塊..."):
                text_chunks = split_text_into_chunks(extracted_text, MAX_CHARS_PER_CHUNK)
                total_chunks = len(text_chunks)
            if total_chunks == 0: st.warning("未能將文本有效分割成塊。"); st.stop()
            st.info(f"文本已分割成 {total_chunks} 個塊進行初步翻譯。")

            # 步驟 3: 逐塊翻譯
            translated_chunks = []
            errors_in_translation = False
            progress_bar = progress_placeholder.progress(0)
            status_text = progress_placeholder.text(f"正在翻譯塊 1 / {total_chunks}...")
            fixed_target_language = "繁體中文"

            for i, chunk in enumerate(text_chunks):
                chunk_num = i + 1
                status_text.text(f"初步翻譯：塊 {chunk_num} / {total_chunks}...")
                translated_chunk = translate_text(chunk, fixed_target_language)
                if translated_chunk and "[[翻譯錯誤:" not in translated_chunk:
                    translated_chunks.append(translated_chunk)
                else:
                    errors_in_translation = True
                    translated_chunks.append(f"\n--- 塊 {chunk_num} 初步翻譯失敗 ---\n{translated_chunk or '未知錯誤'}\n---")
                    st.error(f"初步翻譯塊 {chunk_num} 時發生錯誤。")
                progress_bar.progress(chunk_num / total_chunks)
                if chunk_num < total_chunks: time.sleep(API_CALL_DELAY)
            status_text.text("初步翻譯完成，正在合併結果...")

            # 步驟 4: 合併初步翻譯結果
            raw_final_translated_text = "\n\n".join(translated_chunks)
            raw_result_placeholder.text_area(
                label="raw_translation_label_updated", value=raw_final_translated_text, height=250,
                key="raw_result_updated", disabled=False, label_visibility="collapsed"
            )
            if errors_in_translation: st.warning("部分文本塊初步翻譯失敗，請檢查結果。")
            else: st.success("所有文本塊初步翻譯完成！")

            # 提供初步翻譯的 Docx 下載
            docx_data_raw, docx_error_raw = create_docx_from_text(raw_final_translated_text, base_filename, "_初步翻譯")
            if not docx_error_raw and docx_data_raw['data']:
                raw_download_placeholder.download_button(
                    label=f"📥 下載初步翻譯 ({docx_data_raw['name']})", data=docx_data_raw['data'],
                    file_name=docx_data_raw['name'], mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    key='download_docx_raw'
                )
            else: raw_download_placeholder.error("產生初步翻譯 Docx 時出錯。")

            # --- 新增步驟 5: 潤飾翻譯結果 ---
            status_text.text("正在進行 AI 潤飾與校對...")
            progress_bar.progress(0) # 可以重設進度條或用新的
            with st.spinner("AI 正在努力潤飾中，請稍候..."): # 全局 spinner
                polished_text = polish_translation(raw_final_translated_text, fixed_target_language)
                progress_bar.progress(1) # 潤飾完成
                status_text.text("AI 潤飾完成！")

            if polished_text and "[[潤飾錯誤:" not in polished_text:
                polished_result_placeholder.text_area(
                    label="polished_translation_label_updated", value=polished_text, height=250,
                    key="polished_result_updated", disabled=False, label_visibility="collapsed"
                )
                st.success("AI 潤飾與校對完成！")
                # 提供潤飾後翻譯的 Docx 下載
                docx_data_polished, docx_error_polished = create_docx_from_text(polished_text, base_filename, "_潤飾版")
                if not docx_error_polished and docx_data_polished['data']:
                    polished_download_placeholder.download_button(
                        label=f"📥 下載潤飾後翻譯 ({docx_data_polished['name']})", data=docx_data_polished['data'],
                        file_name=docx_data_polished['name'], mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        key='download_docx_polished'
                    )
                else: polished_download_placeholder.error("產生潤飾版 Docx 時出錯。")
            else:
                polished_result_placeholder.text_area(
                    label="polished_translation_label_error", value=f"AI 潤飾失敗或無結果。\n{polished_text or '詳細錯誤請看上方訊息。'}", height=250,
                    key="polished_result_error", disabled=False, label_visibility="collapsed"
                )
                st.error("AI 潤飾步驟失敗。")
        # ... (處理提取文字失敗或為空的情況，保持不變) ...
        elif extracted_text is not None and not extracted_text.strip():
             st.warning("從檔案中未提取到任何有效文字內容。")
             raw_result_placeholder.text_area("raw_translation_label_no_text", value="未提取到文字。", height=250, key="raw_no_text", disabled=False, label_visibility="collapsed")
        else:
             st.error("無法從文件中提取文字。")
             raw_result_placeholder.text_area("raw_translation_label_extract_fail", value="文件讀取或文字提取失敗。", height=250, key="raw_extract_fail", disabled=False, label_visibility="collapsed")